import pdfplumber
from docx import Document
import re
import os
import logging
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


class ResumeParser:
    def __init__(self):
        self.skill_keywords = {
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby',
            'php', 'swift', 'kotlin', 'rust', 'golang', 'scala', 'r', 'matlab',

            # Web Frontend
            'react', 'angular', 'vue', 'html', 'css', 'sass', 'tailwind',
            'bootstrap', 'next.js', 'nuxt.js', 'gatsby', 'webpack', 'vite',

            # Web Backend
            'node.js', 'django', 'flask', 'fastapi', 'spring', 'express',
            'laravel', 'rails', 'asp.net', 'graphql', 'rest api', 'microservices',

            # Databases
            'sql', 'nosql', 'mongodb', 'postgresql', 'mysql', 'redis',
            'elasticsearch', 'sqlite', 'cassandra', 'dynamodb', 'firebase',

            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins',
            'terraform', 'ansible', 'nginx', 'linux', 'ci/cd', 'git',
            'github', 'gitlab', 'bitbucket',

            # AI & Data
            'machine learning', 'deep learning', 'ai', 'data science',
            'pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn',
            'computer vision', 'nlp', 'tableau', 'power bi', 'excel',

            # Mobile
            'android', 'ios', 'react native', 'flutter',

            # Soft Skills
            'agile', 'scrum', 'leadership', 'communication',
            'project management', 'problem solving', 'teamwork',

            # Design
            'figma', 'photoshop', 'illustrator', 'ui/ux',
        }

    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF, DOCX or TXT file"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Check file size
        file_size = os.path.getsize(file_path)
        max_size = 10 * 1024 * 1024  # 10MB
        if file_size > max_size:
            raise ValueError(f"File too large: {file_size} bytes. Max allowed: {max_size} bytes")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF with page limit"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                # Limit to first 15 pages
                pages = pdf.pages[:15]
                for page in pages:
                    # Extract regular text
                    t = page.extract_text()
                    if t:
                        text += t + "\n"

                    # Extract text from tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                row_text = " ".join(
                                    [cell for cell in row if cell]
                                )
                                text += row_text + "\n"
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Could not read PDF file: {e}")
        return text

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX including tables"""
        text = ""
        try:
            doc = Document(file_path)

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Could not read DOCX file: {e}")
        return text

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise ValueError(f"Could not read TXT file: {e}")

    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        # Remove multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove special characters but keep common punctuation
        text = re.sub(r'[^\w\s@.,;:()\-+/]', ' ', text)
        return text.strip()

    def parse_resume(self, file_path: str) -> Dict:
        """Main method to parse resume and extract all data"""
        try:
            raw_text = self.extract_text(file_path)
            text = self._clean_text(raw_text)

            return {
                'name': self._extract_name(text),
                'email': self._extract_email(text),
                'phone': self._extract_phone(text),
                'skills': self._extract_skills(text),
                'experience': self._extract_experience(text),
                'education': self._extract_education(text),
                'projects': self._extract_projects(text),
                'certifications': self._extract_certifications(text),
                'full_text': text[:8000]  # Limit stored text to 8000 chars
            }
        except Exception as e:
            logger.error(f"Resume parsing error: {e}")
            raise

    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume text"""
        lines = [
            line.strip()
            for line in text.split('\n')
            if line.strip()
        ]

        # Check first 5 lines for a name-like string
        for line in lines[:5]:
            # Skip lines that look like headers or contact info
            if any(skip in line.lower() for skip in [
                'resume', 'curriculum', 'cv', 'profile',
                '@', 'http', 'www', 'phone', 'email',
                'address', 'linkedin', 'github'
            ]):
                continue

            # Name should be 2-4 words, mostly alphabetic
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(re.match(r'^[A-Za-z][a-zA-Z\-\.]*$', w) for w in words):
                    return line

        return lines[0] if lines else "Candidate"

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        match = re.search(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
            text
        )
        return match.group(0).lower() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number supporting international formats"""
        patterns = [
            r'\+?1?\s?\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}',  # US/CA
            r'\+\d{1,3}[\s.\-]?\d{2,4}[\s.\-]?\d{3,4}[\s.\-]?\d{3,4}',  # International
            r'\d{5}[\s.\-]?\d{5}',  # India style
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0).strip()
        return None

    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from text - public method"""
        return self._extract_skills(text)

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        text_lower = text.lower()
        found = []

        for skill in self.skill_keywords:
            # Use word boundary for short skills to avoid false positives
            if len(skill) <= 3:
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text_lower):
                    found.append(skill)
            else:
                if skill in text_lower:
                    found.append(skill)

        # Sort by relevance (frequency in text)
        found.sort(key=lambda s: text_lower.count(s), reverse=True)
        return found

    def _extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience entries"""
        experience = []

        # Look for date ranges like "Jan 2020 - Present" or "2018 - 2021"
        date_pattern = re.compile(
            r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?'
            r'\.?\s*\d{4})\s*[-–to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|'
            r'Jul|Aug|Sep|Oct|Nov|Dec)?\.?\s*\d{4}|Present|Current|Now)',
            re.IGNORECASE
        )

        matches = date_pattern.finditer(text)
        for match in matches:
            start = match.group(1).strip()
            end = match.group(2).strip()

            # Get surrounding context (the line with the date)
            start_pos = max(0, match.start() - 100)
            end_pos = min(len(text), match.end() + 100)
            context = text[start_pos:end_pos].strip()
            context = re.sub(r'\s+', ' ', context)

            experience.append({
                'start_date': start,
                'end_date': end,
                'description': context
            })

        # Also look for years of experience mentions
        years_match = re.findall(
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            text,
            re.IGNORECASE
        )
        if years_match and not experience:
            experience.append({
                'years': int(years_match[0]),
                'description': f"{years_match[0]} years of experience"
            })

        return experience[:10]  # Limit to 10 entries

    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education entries"""
        keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'b.s.', 'm.s.',
            'b.tech', 'm.tech', 'b.e.', 'm.e.', 'mba', 'university',
            'college', 'institute', 'school of'
        ]
        results = []

        for line in text.split('\n'):
            line_lower = line.lower()
            if any(kw in line_lower for kw in keywords):
                clean = line.strip()
                if clean and 10 < len(clean) < 200:
                    # Try to extract year
                    year_match = re.search(r'\b(19|20)\d{2}\b', clean)
                    results.append({
                        'description': clean,
                        'year': year_match.group(0) if year_match else None
                    })

        return results[:5]  # Limit to 5 entries

    def _extract_projects(self, text: str) -> List[str]:
        """Extract project descriptions"""
        keywords = [
            'project', 'developed', 'built', 'created',
            'designed', 'implemented', 'architected'
        ]
        results = []

        # Split by newline instead of period to avoid .js issues
        for line in text.split('\n'):
            line = line.strip()
            if (
                any(kw in line.lower() for kw in keywords)
                and 20 < len(line) < 300
            ):
                results.append(line)

        return results[:5]  # Limit to 5 entries

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certification entries"""
        keywords = [
            'certified', 'certification', 'certificate',
            'coursera', 'udemy', 'aws certified', 'google certified',
            'microsoft certified', 'comptia', 'cisco', 'pmp', 'cpa'
        ]
        results = []

        for line in text.split('\n'):
            line = line.strip()
            if (
                any(kw in line.lower() for kw in keywords)
                and 10 < len(line) < 200
            ):
                results.append(line)

        return results[:5]  # Limit to 5 entries